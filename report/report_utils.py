# report_utils.py
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

import plotly.io as pio

from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
from plotly.graph_objects import Figure

import uuid
import tempfile
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager



class DocxReportUtils:
    def __init__(self, template_path: str):
        self.doc = Document(template_path)

    def insert_table(self, placeholder: str, columns: list, data: list, col_widths_percentage: list = None,
                     alt_row_color="D6E9F2"):
        for para in self.doc.paragraphs:
            if placeholder in para.text:
                table = self.doc.add_table(rows=1, cols=len(columns))
                table.autofit = False

                section = self.doc.sections[0]
                total_width = section.page_width - section.left_margin - section.right_margin

                # Calculate column widths from percentages if provided
                if col_widths_percentage:
                    if not abs(sum(col_widths_percentage) - 1.0) < 1e-6:
                        raise ValueError("Column widths percentages must sum to 1.0")
                    col_widths = [total_width * p for p in col_widths_percentage]
                else:
                    col_widths = [total_width / len(columns)] * len(columns)

                hdr_cells = table.rows[0].cells
                for idx, col_name in enumerate(columns):
                    hdr_cells[idx].text = col_name
                    hdr_cells[idx].width = col_widths[idx]
                    self._format_cell(hdr_cells[idx], bold=True)

                self._add_border(table.rows[0], "top", "8")
                self._add_border(table.rows[0], "bottom", "12")
                self._add_border(table.rows[-1], "bottom", "12")

                for i, row_data in enumerate(data):
                    row_cells = table.add_row().cells
                    for j, value in enumerate(row_data):
                        row_cells[j].text = str(value)
                        row_cells[j].width = col_widths[j]
                        bg = alt_row_color if i % 2 == 0 else "FFFFFF"  # Alternating colors
                        self._format_cell(row_cells[j], bg_color=bg)

                para.text = ''
                para._element.addnext(table._element)
                break

    def _add_border(self, row, position=None, thickness: str = "8"):
        supported_values = ["bottom", "top", "right", "left"]
        if position not in supported_values:
            raise ValueError(f"'position' supported values are {supported_values}")
        elif position is None:
            position = "bottom"

        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders_list = OxmlElement('w:tcBorders')
            modified_border = OxmlElement(f'w:{position}')
            modified_border.set(qn('w:val'), 'single')
            modified_border.set(qn('w:sz'), thickness)  # thickness
            modified_border.set(qn('w:color'), '000000')  # black
            borders_list.append(modified_border)
            tcPr.append(borders_list)

    def insert_image(self, placeholder: str, plotly_fig, width_px=1200, height_px=1000):
        import tempfile
        import uuid

        # Ensure we're always working with a list
        figures = plotly_fig if isinstance(plotly_fig, list) else [plotly_fig]

        section = self.doc.sections[0]
        max_width = section.page_width - section.left_margin - section.right_margin

        for para in self.doc.paragraphs:
            if placeholder in para.text:
                # Clear the placeholder text
                para.text = ''
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                for fig in figures:
                    temp_dir = tempfile.gettempdir()
                    temp_img_path = os.path.join(temp_dir, f"{uuid.uuid4()}.png")
                    # fig.update_layout(margin=dict(l=20, r=20, t=40, b=40))

                    temp_img_path = self.export_fig_as_png_from_html(fig, width_px=width_px, height_px=height_px, output_path=temp_img_path)

                    # Insert image
                    run = para.add_run()
                    run.add_picture(temp_img_path, width=max_width)

                    # Add a line break after each image except the last one
                    if fig != figures[-1]:
                        para.add_run().add_break()

                    os.remove(temp_img_path)
                break
    @staticmethod
    def get_unique_filename(path):
        """
        If 'path' exists, appends _1, _2, etc. before the file extension.
        """
        base, ext = os.path.splitext(path)
        counter = 1
        new_path = path

        while os.path.exists(new_path):
            new_path = f"{base}_{counter}{ext}"
            counter += 1

        return new_path

    def replace_text(self, placeholder: str, text: str):
        # Replace in main body
        for para in self.doc.paragraphs:
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, text)

        # Replace in headers
        for section in self.doc.sections:
            self._replace_in_header_footer(section.header, placeholder, text)

    def _replace_in_header_footer(self, header_footer, placeholder, text):
        # Check paragraphs
        for para in header_footer.paragraphs:
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, text)

        # Check tables inside header/footer
        for table in header_footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if placeholder in para.text:
                            para.text = para.text.replace(placeholder, text)

    def save(self, output_path: str):
        unique_path = self.get_unique_filename(output_path)
        self.doc.save(f"{unique_path}")

    @staticmethod
    def _format_cell(cell, bold=False, bg_color=None, font_config=None):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                font = run.font
                if font_config:
                    for attr, value in font_config.items():
                        if attr == "size":
                            setattr(font, attr, Pt(value))
                        elif attr == "color":
                            rgb = value if isinstance(value, RGBColor) else RGBColor(*value)
                            font.color.rgb = rgb
                        else:
                            setattr(font, attr, value)
                font.bold = bold
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        if bg_color:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), bg_color)
            cell._tc.get_or_add_tcPr().append(shading)



    @staticmethod
    def export_fig_as_png_from_html(fig: Figure, width_px: int = 150, height_px: int = 150, output_path: str = None) -> str:
        """
        Renders a Plotly figure to HTML, opens it headlessly in Chrome, and exports it as a PNG.
        TODO:
            - Replace this method with Plotly's `to_image()` if Kaleido becomes stable
              in newer Plotly versions and configuration. Current versions: Plotly 6.1.2 and Kaleido 0.2.1.

        """

        html_str = pio.to_html(fig, include_plotlyjs='cdn')
        html_str = html_str.replace(
            "</head>",
            """
            <style>
                body { margin: 0; overflow: hidden; }
            </style>
            </head>
            """
        )

        temp_dir = tempfile.gettempdir()
        html_path = os.path.join(temp_dir, f"{uuid.uuid4()}.html")

        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_str)

        # Set up headless Chrome with correct service usage
        options = Options()
        options.add_argument("--headless=new")
        options.add_argument("--disable-gpu")
        # options.add_argument(f"--window-size={width_px},{height_px}")

        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)

        # Force exact viewport - guarantees pixel-perfect screenshots
        driver.execute_cdp_cmd("Emulation.setDeviceMetricsOverride", {
            "mobile": False,
            "width": width_px,
            "height": height_px,
            "deviceScaleFactor": 1,
        })

        driver.set_window_size(width_px, height_px)
        try:
            driver.get("file://" + html_path)

            # Wait until Plotly's main SVG is rendered
            WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.CLASS_NAME, "main-svg"))
            )

            driver.save_screenshot(output_path)
        finally:
            if os.path.exists(html_path):
                os.remove(html_path)
            driver.quit()

        return output_path
